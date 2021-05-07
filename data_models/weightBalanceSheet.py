# -*- coding: utf-8 -*-

# 开发者：许鑫
# 时间：2020/9/5-2020/9/13
# 功能：自动生成重量平衡表word
import os

from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from data_models.data_collector import aircraft as af

#
# class WeightBalanceSheet(object):
#     def __init__(self):
#         name = ''
#
#     def report(self):


def export_wab_sheet(save_dir, aircraft, wab_id, wab_version, task_id, weigh_basis, als_basis,
                     test_type, crew_num, date):
    aircraft_weight_info = af.get_aircraft_weight_info()
    records = {'机型/架机：': aircraft,
               '重量平衡表编号：': wab_id,
               '重量平衡表版本：': wab_version}
    records2 = {'任务单编号：': task_id,
                '试验空机重量重心依据：': weigh_basis,
                '配载数据依据：': als_basis}
    records3 = {'试验内容：': test_type,
                '上机人数': crew_num,
                '编制日期：': date}
    test_empty_weight = aircraft_weight_info['major_aircraft_weight']['test_empty_weight']
    records4 = {'试验空机重量(W)': {'说明/备注': '', '重量(kg)': str(test_empty_weight[1]),
                              '力臂(mm )': str(test_empty_weight[2]),
                              '力矩(kg*mm)': str(test_empty_weight[3]),
                              '重心': str(test_empty_weight[4])}}
    for item in aircraft_weight_info['operation_items']:
        records4[item[0]] = {'说明/备注': '', '重量(kg)': str(item[1]), '力臂(mm )': str(item[2]),
                             '力矩(kg*mm)': str(item[3]), '重心': ''}
    # records4 = {'试验空机重量(W)':
    #                 {'说明/备注': '', '重量(kg)': '46349', '力臂(mm )': '20593', '力矩(kg*mm)': '954492117', '重心': '17.7'},
    #             '驾驶员':
    #                 {'说明/备注': '2×75 kg', '重量(kg)': '150', '力臂(mm )': '5990', '力矩(kg*mm)': '898500', '重心': ''},
    #             '观察员':
    #                 {'说明/备注': '75 kg', '重量(kg)': '75', '力臂(mm )': '6900', '力矩(kg*mm)': '517500', '重心': ''},
    #             '试飞工程师A、B ':
    #                 {'说明/备注': '2×75 kg', '重量(kg)': '150', '力臂(mm )': '21442', '力矩(kg*mm)': '3216270', '重心': ''},
    #             '试飞工程师C':
    #                 {'说明/备注': '75 kg', '重量(kg)': '75', '力臂(mm )': '21459', '力矩(kg*mm)': '1609448', '重心': ''},
    #             '机组行李':
    #                 {'说明/备注': '', '重量(kg)': '60', '力臂(mm )': '8417', '力矩(kg*mm)': '505020', '重心': ''}
    #             }
    records5 = {'使用空重(OEW)':
                    {'说明/备注': '', '重量(kg)': '46859', '力臂(mm )': '20513', '力矩(kg*mm)': '961238854', '重心': '15.8'},
                '驾驶员':
                    {'说明/备注': '', '重量(kg)': '7000', '力臂(mm )': '23835', '力矩(kg*mm)': '166845700', '重心': ''},
                '观察员':
                    {'说明/备注': '', '重量(kg)': '3527', '力臂(mm )': '17586', '力矩(kg*mm)': '62026464', '重心': ''}
                }

    ###加底纹方法
    def add_shading(hdr_cells):
        shading_elm = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr_cells._tc.get_or_add_tcPr().append(shading_elm)

    ###--begin--设置文档类型(字体+字号)
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = u'宋体'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    font.name = u'Times New Roman'
    font.size = Pt(9)
    ###begin--页边距的设置
    sec = document.sections[0]  # sections对应文档中的“节”
    sec.left_margin = Inches(0.5874)  # 以下依次设置左、右、上、下页面边距
    sec.right_margin = Inches(0.5874)
    sec.top_margin = Inches(0.5874)
    sec.bottom_margin = Inches(0.4874)
    sec.page_width = Inches(8.27)  # 设置页面宽度
    sec.page_height = Inches(11.69)  # 设置页面高度
    ###end
    ###begin--设置页眉、页脚
    header = document.sections[0].header  # 获取第一个节的页眉
    p = header.paragraphs[0]  # 获取页眉的第一个段落
    p_header = p.add_run('')
    # p_header.add_picture('monty-truth.png', width=Inches(1.35))  # 添加页眉内容
    p_header_text_blank = p.add_run('                                                           ')
    p_header_text = p.add_run('FM8007-003D')

    footer = document.sections[0].footer  # 获取第一个节的页脚
    p = footer.paragraphs[0]  # 获取页脚的第一个段落
    p_footer = p.add_run('')  # 获取页脚的第一个段落
    # p_footer.add_picture('footer.png', width=Inches(7.1))  # 添加页脚内容
    ###begin--表格标题
    p = document.add_paragraph()
    p_text = p.add_run(u'重量平衡表')
    p_text.font.bold = True
    # p_text.font.name = u'宋体'
    # p_text._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p_text.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ###设置段前段后距
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    ###end

    ###begin--创建表格1
    table1 = document.add_table(rows=1, cols=1, style='Table Grid')
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '1)重量平衡计算表信息'
    ###设置表格格式
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER  ###居中
    hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###垂直居中
    table1.rows[0].height = Cm(0.52)  ###设置行高
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True  ###设置粗体
    ###加底纹
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
    hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    ###end

    ###创建表格2
    ###begin--设置表格内容
    table2 = document.add_table(rows=3, cols=7, style='Table Grid')
    hdr2_col0_cells = table2.columns[0].cells
    hdr2_col0_cells[0].text = '机型/架机：'
    hdr2_col0_cells[1].text = '重量平衡表编号：'
    hdr2_col0_cells[2].text = '重量平衡表版本：'
    hdr2_col2_cells = table2.columns[2].cells
    hdr2_col2_cells[0].text = '任务单编号：'
    hdr2_col2_cells[1].text = '试验空机重量重心依据：'
    hdr2_col2_cells[2].text = '配载数据依据：'
    hdr2_col5_cells = table2.columns[5].cells
    hdr2_col5_cells[0].text = '试验内容：'
    hdr2_col5_cells[1].text = '上机人数：'
    hdr2_col5_cells[2].text = '编制日期：'
    hdr2_col1_cells = table2.columns[1].cells
    hdr2_col3_cells = table2.columns[3].cells
    hdr2_col4_cells = table2.columns[4].cells
    hdr2_col5_cells = table2.columns[5].cells
    hdr2_col6_cells = table2.columns[6].cells
    ###--end
    ###为表格添加内容
    for i, value, value2, value3 in zip(range(len(table2.rows)), records.values(), records2.values(),
                                        records3.values()):
        hdr2_col1_cells[i].text = value
        hdr2_col4_cells[i].text = value2
        hdr2_col6_cells[i].text = value3
    ###begin--设置表格格式
    # table2.autofit = False ###设置为手动调整表格宽度
    for i in range(3):
        table2.rows[i].height = Cm(0.52)  ###设置行高
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER  ###居中显示
    ###合并单元格水平
    for i in range(3):
        hdr_cells = table2.rows[i].cells
        hdr_cells[2].merge(hdr_cells[3])
    ###设置列宽
    for cell in hdr2_col0_cells:
        cell.width = Inches(0.2)
    for cell in hdr2_col1_cells:
        cell.width = Inches(1.9)
    for cell in hdr2_col2_cells:
        cell.width = Inches(2.1)
    for cell in hdr2_col4_cells:
        cell.width = Inches(1.8)
    for cell in hdr2_col5_cells:
        cell.width = Inches(0.9)
    # for cell in hdr2_col6_cells:
    #     cell.width = Inches(0.6)
    ###加底纹
    for i in range(len(table2.rows)):
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr2_col0_cells[i]._tc.get_or_add_tcPr().append(shading_elm_2)
    for i in range(len(table2.rows)):
        shading_elm_3 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr2_col2_cells[i]._tc.get_or_add_tcPr().append(shading_elm_3)
    for i in range(len(table2.rows)):
        shading_elm_4 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr2_col5_cells[i]._tc.get_or_add_tcPr().append(shading_elm_4)
    ###水平居中和垂直居中
    for i in range(len(table2.rows)):
        for j in range(len(table2.columns)):
            table2.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  ###方法二水平居中
            table2.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###方法二垂直居中
    ###--end

    ###begin--创建表格3
    table3 = document.add_table(rows=1, cols=1, style='Table Grid')
    hdr3_row0_cells = table3.rows[0].cells
    hdr3_row0_cells[0].text = '2)试验机各项目重量平衡信息'
    ###设置表格格式
    table3.rows[0].height = Cm(0.52)  ###设置行高
    hdr3_row0_cells[0].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER  ###居中
    hdr3_row0_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###垂直居中
    ###begin--加底纹
    for i in range(len(table3.rows)):
        shading_elm_5 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr3_row0_cells[i]._tc.get_or_add_tcPr().append(shading_elm_5)
    ###--end

    ###begin--创建表格4
    table4 = document.add_table(rows=20, cols=8, style='Table Grid')
    table4.autofit = False  ###设置为手动调整表格宽度
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER  ###居中显示
    ###设置表格内容
    hdr4_row0_cells = table4.rows[0].cells
    hdr4_row0_cells[0].text = '项目'
    hdr4_row0_cells[1].text = '说明/备注'
    hdr4_row0_cells[2].text = '重量(kg)'
    hdr4_row0_cells[4].text = '力臂(mm)'
    hdr4_row0_cells[5].text = '力矩(kg*mm)'
    hdr4_row0_cells[6].text = '重心(%MAC)'
    hdr4_row0_cells[7].text = '配平(°)'
    table4.rows[0].height = Cm(0.87)
    hdr4_row16_cells = table4.rows[16].cells
    hdr4_row16_cells[4].text = '\\'
    hdr4_row16_cells[5].text = '\\'
    hdr4_row17_cells = table4.rows[17].cells
    hdr4_row17_cells[4].text = '\\'
    hdr4_row17_cells[5].text = '\\'

    hdr4_col0_cells = table4.columns[0].cells
    hdr4_col0_cells[1].text = '试验空机重量(W)'

    hdr4_col0_cells[2].text = '驾驶员'
    hdr4_col0_cells[3].text = '观察员'
    hdr4_col0_cells[4].text = '试飞工程师A、B'
    hdr4_col0_cells[5].text = '试飞工程师C'
    hdr4_col0_cells[6].text = '机组行李'
    hdr4_col0_cells[7].text = ''
    hdr4_col0_cells[8].text = ''
    hdr4_col0_cells[9].text = ''
    hdr4_col0_cells[10].text = '使用空重(OEW)'
    hdr4_col0_cells[11].text = '试验固定配重'
    hdr4_col0_cells[12].text = '水配重'
    hdr4_col0_cells[13].text = '零油重量(ZFW)'
    hdr4_col0_cells[14].text = '燃油总量'
    hdr4_col0_cells[16].text = '左/右油箱油量'
    hdr4_col0_cells[17].text = '中央翼油箱油量'
    hdr4_col0_cells[18].text = '起飞重量(TOW)'

    hdr4_col7_cells = table4.columns[7].cells
    hdr4_col7_cells[1].text = '\\'

    hdr4_col6_cells = table4.columns[6].cells
    hdr4_col5_cells = table4.columns[5].cells
    hdr4_col2_cells = table4.columns[2].cells
    hdr4_col6_cells[2].text = '\\'
    hdr4_col6_cells[11].text = '\\'
    hdr4_col6_cells[12].text = '\\'
    hdr4_col6_cells[14].text = '\\'
    ###为表格添加内容
    # for temp_key in records4.keys():
    for i, temp_key in zip(range(1, 7), records4.keys()):
        for j, value in zip(range(1, 7), records4[temp_key].values()):
            if value == '':
                continue
            if j >= 3:
                table4.cell(i, j + 1).text = value  ###跳过合并的第3列
            else:
                table4.cell(i, j).text = value

    for i, temp_key in zip(range(10, 13), records5.keys()):
        for j, value in zip(range(1, 7), records5[temp_key].values()):
            if value == '':
                continue
            if j >= 3:
                table4.cell(i, j + 1).text = value  ###跳过合并的第3列
            else:
                table4.cell(i, j).text = value
    ###设置表格格式
    for i in range(20):
        table4.rows[i].height = Cm(0.52)  ###设置行高
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr4_col7_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###垂直居中
    hdr4_col7_cells[1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  ###水平居中
    ###设置字体粗体
    hdr4_col0_cells[1].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col0_cells[10].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col0_cells[13].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col0_cells[18].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col2_cells[1].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col2_cells[10].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col6_cells[1].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col6_cells[10].paragraphs[0].runs[0].font.bold = True  ###设置为粗体

    for i in range(len(table4.columns)):
        if i != 3:
            hdr4_row0_cells[i].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr4_col0_cells[18].style = "Times New Roman"  ###设置为粗体
    ###加底纹
    for i in range(len(table4.columns)):
        shading_elm_6 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr4_row0_cells[i]._tc.get_or_add_tcPr().append(shading_elm_6)
    for i in range(len(table4.rows)):
        shading_elm_6 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr4_col0_cells[i]._tc.get_or_add_tcPr().append(shading_elm_6)
    hdr4_col2_cells = table4.columns[2].cells
    for i in range(len(table4.rows)):
        shading_elm_7 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr4_col2_cells[i]._tc.get_or_add_tcPr().append(shading_elm_7)
    hdr4_col3_cells = table4.columns[3].cells
    for i in range(len(table4.rows)):
        shading_elm_8 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
        hdr4_col3_cells[i]._tc.get_or_add_tcPr().append(shading_elm_8)
    add_shading(hdr4_col6_cells[10])
    add_shading(hdr4_col6_cells[1])
    add_shading(hdr4_col6_cells[13])
    add_shading(hdr4_col6_cells[18])
    add_shading(hdr4_col7_cells[18])
    add_shading(hdr4_col7_cells[19])
    ###合并单元格垂直
    for i in range(6):
        hdr_cells = table4.columns[i].cells
        hdr_cells[18].merge(hdr_cells[19])

    hdr4_col6_cells[2].merge(hdr4_col6_cells[9])
    hdr4_col6_cells[14].merge(hdr4_col6_cells[17])
    hdr4_col6_cells[18].merge(hdr4_col6_cells[19])

    hdr4_col7_cells[1].merge(hdr4_col7_cells[17])

    for i in range(6):
        hdr_cells = table4.columns[i].cells
        if i != 2 and i != 3:
            hdr_cells[14].merge(hdr_cells[15])
    ###合并单元格水平
    for i in range(14):
        hdr_cells = table4.rows[i].cells
        hdr_cells[2].merge(hdr_cells[3])

    hdr4_row18_cells = table4.rows[18].cells
    hdr4_row18_cells[2].merge(hdr4_row18_cells[3])

    hdr4_row0_cells = table4.rows[0].cells
    hdr4_row0_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###垂直居中
    ###设置列宽
    for cell in hdr4_col0_cells:
        cell.width = Inches(1.6)
    hdr4_col1_cells = table4.columns[1].cells
    for cell in hdr4_col1_cells:
        cell.width = Inches(1)
    hdr4_col4_cells = table4.columns[4].cells
    for cell in hdr4_col4_cells:
        cell.width = Inches(1.1)
    hdr4_col5_cells = table4.columns[5].cells
    for cell in hdr4_col5_cells:
        cell.width = Inches(1.2)
    hdr4_col7_cells = table4.columns[7].cells
    for cell in hdr4_col7_cells:
        cell.width = Inches(0.9)
    ###水平居中和垂直居中
    for i in range(len(table4.rows)):
        for j in range(len(table4.columns)):
            table4.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  ###方法二水平居中
            table4.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###方法二垂直居中
    ###end

    ###begin--创建表格5
    table5 = document.add_table(rows=6, cols=1, style='Table Grid')
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr5_col0_cells = table5.columns[0].cells
    hdr5_col0_cells[0].text = '3)燃油消耗曲线'
    ###设置表格格式
    for i in range(0, 6):
        hdr5_col0_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###垂直居中
        table5.rows[i].height = Cm(0.52)  ###设置行高
    # table5.rows[3].height = Cm(0.52)  ###设置行高
    # hdr5_col0_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ###垂直居中
    ###begin--插入燃油消耗曲线图片
    hdr5_col0_cells[0].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    p = hdr5_col0_cells[1].add_paragraph()
    # p.add_run().add_picture("fuel_consumption.png", height=Inches(1.9))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ###--end
    hdr5_col0_cells[2].text = '4)签署'
    hdr5_col0_cells[2].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr5_col0_cells[
        3].text = '编制:                                     校对:                                     审核:                                     会签:                                     批准:                                     '
    hdr5_col0_cells[4].text = '5)变更记录'
    hdr5_col0_cells[4].paragraphs[0].runs[0].font.bold = True  ###设置为粗体
    hdr5_col0_cells[
        5].text = '编制:                                     会签:                                     审核:      '
    table5.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    table5.rows[1].height = Cm(5.45)
    table5.rows[5].height = Cm(0.8)
    ###begin--加底纹
    shading_elm_10 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
    hdr5_col0_cells[0]._tc.get_or_add_tcPr().append(shading_elm_10)
    shading_elm_11 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
    hdr5_col0_cells[2]._tc.get_or_add_tcPr().append(shading_elm_11)
    shading_elm_12 = parse_xml(r'<w:shd {} w:fill="#EEECE1"/>'.format(nsdecls('w')))
    hdr5_col0_cells[4]._tc.get_or_add_tcPr().append(shading_elm_12)
    ###end
    # document.add_page_break()
    try:
        document.save(save_dir + os.sep + '重量平衡表' + wab_id + '.docx')
    except PermissionError:
        return '导出重量平衡表失败：无文件写入权限！'
    return None
